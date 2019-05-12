from docx import Document
from docx.shared import Cm, Pt
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

import os, shutil


def makeCertification(**context):
    document = Document()

    table = document.add_table(rows=20, cols=8)
    table.style = "TableGrid"

    # table size
    widthList = [1, 1, 1, 2.5, 2.5, 2.5, 2.5, 2.5]
    for row in table.rows:
        row.height = Cm(0.8)
        for col, width in zip(row.cells, widthList):
            col.width = Cm(width)

    # merge cells
    pictureStart = table.cell(0, 0)
    pictureEnd = table.cell(4, 2)
    picture = pictureStart.merge(pictureEnd)
    if context.get('image', "") != "":
        pictureRun = picture.paragraphs[0].add_run()
        pictureRun.add_picture(context['image'], width=Cm(3), height=Cm(3))
    else:
        pictureRun = picture.paragraphs[0].add_run("사진")


    titleStart = table.cell(0, 3)
    titleEnd = table.cell(1, 7)
    title = titleStart.merge(titleEnd)
    titleRun = title.paragraphs[0].add_run("이력서")
    titleRun.font.size = Pt(20)
    titleRun.bold = True

    nameStart = table.cell(2, 3)
    nameEnd = table.cell(3, 3)
    name = nameStart.merge(nameEnd)
    nameRun = name.paragraphs[0].add_run("이름")
    nameRun.bold =True

    nameContentStart = table.cell(2, 4)
    nameContentEnd = table.cell(3, 5)
    nameContent = nameContentStart.merge(nameContentEnd)
    nameContent.text = context.get("name", "")

    sex = table.cell(2, 6)
    sexRun = sex.paragraphs[0].add_run("성별")
    sexRun.bold =True
    sexContent = table.cell(3, 6)
    sexContent.text = context.get("sex", "")

    age = table.cell(2, 7)
    ageRun = age.paragraphs[0].add_run("나이")
    ageRun.bold =True
    ageContent = table.cell(3, 7)
    ageContent.text = context.get("age", "")

    addressStart = table.cell(5, 0)
    addressEnd = table.cell(5, 2)
    address = addressStart.merge(addressEnd)
    addressRun = address.paragraphs[0].add_run("주소")
    addressRun.bold =True

    addressContentStart = table.cell(5, 3)
    addressContentEnd = table.cell(5, 7)
    addressContent = addressContentStart.merge(addressContentEnd)
    addressContent.text = context.get("address", "")

    emailStart = table.cell(6, 0)
    emailEnd = table.cell(6, 2)
    email = emailStart.merge(emailEnd)
    emailRun = email.paragraphs[0].add_run("e-mail")
    emailRun.bold =True

    highNameStart = table.cell(7, 0)
    highNameEnd = table.cell(7, 2)
    highName = highNameStart.merge(highNameEnd)
    highNameRun = highName.paragraphs[0].add_run("고등학교 이름")
    highNameRun.bold =True

    highNameContent = table.cell(7, 3)
    highNameContent.text = context.get("highSchoolName", "")

    highAdDate = table.cell(7, 4)
    highAdDateRun = highAdDate.paragraphs[0].add_run("입학일")
    highAdDateRun.bold =True

    highAdDateContent = table.cell(7, 5)
    highAdDateContent.text = context.get("highSchoolAdmDate", "")

    highGrDate = table.cell(7, 6)
    highGrDateRun = highGrDate.paragraphs[0].add_run("졸업일")
    highGrDateRun.bold =True

    highGrDateContent = table.cell(7, 7)
    highGrDateContent.text = context.get("highSchoolGradDate", "")

    univNameStart = table.cell(8, 0)
    univNameEnd = table.cell(8, 2)
    univName = univNameStart.merge(univNameEnd)
    univNameRun = univName.paragraphs[0].add_run("대학교 이름")
    univNameRun.bold =True

    univNameContent = table.cell(8, 3)
    univNameContent.text = context.get("univSchoolName", "")

    univAdDate = table.cell(8, 4)
    univAdDateRun = univAdDate.paragraphs[0].add_run("입학일")
    univAdDateRun.bold =True

    univAdDateContent = table.cell(8, 5)
    univAdDateContent.text = context.get("univSchoolAdmDate", "")

    univGrDate = table.cell(8, 6)
    univGrDateRun = univGrDate.paragraphs[0].add_run("졸업일")
    univGrDateRun.bold =True

    univGrDateContent = table.cell(8, 7)
    univGrDateContent.text = context.get("univSchoolGradDate", "")

    compTitleStart = table.cell(9, 0)
    compTitleEnd = table.cell(9, 7)
    compTitle = compTitleStart.merge(compTitleEnd)
    compTitleRun = compTitle.paragraphs[0].add_run("경력사항")
    compTitleRun.font.size = Pt(15)
    compTitleRun.bold =True

    compDateSStart = table.cell(10, 0)
    compDateSEnd = table.cell(10, 1)
    compDateS = compDateSStart.merge(compDateSEnd)
    compDateSRun = compDateS.paragraphs[0].add_run("근무시작")
    compDateSRun.bold =True

    compDateEStart = table.cell(10, 2)
    compDateEEnd = table.cell(10, 3)
    compDateE = compDateEStart.merge(compDateEEnd)
    compDateERun = compDateE.paragraphs[0].add_run("근무종료")
    compDateERun.bold =True

    compNameStart = table.cell(10, 4)
    compNameEnd = table.cell(10, 7)
    compName = compNameStart.merge(compNameEnd)
    compNameRun = compName.paragraphs[0].add_run("회사명/근무지")
    compNameRun.bold =True

    compDateSContentStart = table.cell(11, 0)
    compDateSContentEnd = table.cell(11, 1)
    compDateSContent = compDateSContentStart.merge(compDateSContentEnd)
    compDateSContent.text = context.get("compFirstDate", "")

    compDateEContentStart = table.cell(11, 2)
    compDateEContentEnd = table.cell(11, 3)
    compDateEContent = compDateEContentStart.merge(compDateEContentEnd)
    compDateEContent.text = context.get("compLastDate", "")

    compNameContentStart = table.cell(11, 4)
    compNameContentEnd = table.cell(11, 7)
    compNameContent = compNameContentStart.merge(compNameContentEnd)
    compNameContent.text = context.get("compName", "")

    compPlaceContent = table.cell(11, 7)
    
    certDateStart = table.cell(14, 0)
    certDateEnd = table.cell(14, 2)
    certDate = certDateStart.merge(certDateEnd)
    certDateRun = certDate.paragraphs[0].add_run("취득일(연/월/일)")
    certDateRun.bold =True

    certNameStart = table.cell(14, 3)
    certNameEnd = table.cell(14, 6)
    certName = certNameStart.merge(certNameEnd)
    certNameRun = certName.paragraphs[0].add_run("자격증")
    certNameRun.bold =True

    certPlace = table.cell(14, 7)
    certPlaceRun = certPlace.paragraphs[0].add_run("비고")
    certPlaceRun.bold =True


    certDateContentStart = table.cell(15, 0)
    certDateContentEnd = table.cell(15, 2)
    certDateContent = certDateContentStart.merge(certDateContentEnd)
    certDateContent.text = context.get("certGetDate", "")

    certNameContentStart = table.cell(15, 3)
    certNameContentEnd = table.cell(15, 6)
    certNameContent = certNameContentStart.merge(certNameContentEnd)
    certNameContent.text = context.get("certName", "")

    certContent = table.cell(15, 7)

    # align center
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
    return document

if  __name__=="__main__":
    makeCertification()
    